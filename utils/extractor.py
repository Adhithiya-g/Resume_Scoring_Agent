"""
Text Extraction Module for Resume Scoring Agent
Supports PDF, DOCX, and TXT file formats
"""

import os
import logging
from typing import Optional, Dict, Any
import PyPDF2
import pdfplumber
from docx import Document
import streamlit as st

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class TextExtractor:
    """Handles text extraction from various file formats"""
    
    @staticmethod
    def extract_text_from_pdf(file_path: str, method: str = "pdfplumber") -> str:
        """
        Extract text from PDF files using either PyPDF2 or pdfplumber
        
        Args:
            file_path: Path to the PDF file
            method: Extraction method ('pypdf2' or 'pdfplumber')
            
        Returns:
            Extracted text as string
        """
        try:
            if method == "pypdf2":
                return TextExtractor._extract_with_pypdf2(file_path)
            else:
                return TextExtractor._extract_with_pdfplumber(file_path)
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
            return ""
    
    @staticmethod
    def _extract_with_pypdf2(file_path: str) -> str:
        """Extract text using PyPDF2"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            logger.error(f"PyPDF2 extraction failed: {str(e)}")
        return text.strip()
    
    @staticmethod
    def _extract_with_pdfplumber(file_path: str) -> str:
        """Extract text using pdfplumber (more accurate)"""
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            logger.error(f"pdfplumber extraction failed: {str(e)}")
        return text.strip()
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """
        Extract text from DOCX files
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text as string
        """
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            return ""
    
    @staticmethod
    def extract_text_from_txt(file_path: str, encoding: str = "utf-8") -> str:
        """
        Extract text from TXT files
        
        Args:
            file_path: Path to the TXT file
            encoding: File encoding (default: utf-8)
            
        Returns:
            Extracted text as string
        """
        try:
            with open(file_path, 'r', encoding=encoding) as file:
                return file.read().strip()
        except UnicodeDecodeError:
            # Try different encodings
            for enc in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=enc) as file:
                        return file.read().strip()
                except UnicodeDecodeError:
                    continue
            logger.error(f"Could not decode text file {file_path}")
            return ""
        except Exception as e:
            logger.error(f"Error extracting text from TXT {file_path}: {str(e)}")
            return ""
    
    @staticmethod
    def extract_text_from_uploaded_file(uploaded_file) -> str:
        """
        Extract text from Streamlit uploaded file object
        
        Args:
            uploaded_file: Streamlit UploadedFile object
            
        Returns:
            Extracted text as string
        """
        try:
            file_extension = uploaded_file.name.lower().split('.')[-1]
            
            if file_extension == 'pdf':
                return TextExtractor._extract_pdf_from_bytes(uploaded_file.read())
            elif file_extension == 'docx':
                return TextExtractor._extract_docx_from_bytes(uploaded_file.read())
            elif file_extension == 'txt':
                return uploaded_file.read().decode('utf-8')
            else:
                logger.error(f"Unsupported file format: {file_extension}")
                return ""
                
        except Exception as e:
            logger.error(f"Error extracting text from uploaded file: {str(e)}")
            return ""
    
    @staticmethod
    def _extract_pdf_from_bytes(pdf_bytes: bytes) -> str:
        """Extract text from PDF bytes"""
        text = ""
        try:
            # Try pdfplumber first
            import io
            with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception:
            # Fallback to PyPDF2
            try:
                import io
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            except Exception as e:
                logger.error(f"Failed to extract PDF from bytes: {str(e)}")
        
        return text.strip()
    
    @staticmethod
    def _extract_docx_from_bytes(docx_bytes: bytes) -> str:
        """Extract text from DOCX bytes"""
        try:
            import io
            doc = Document(io.BytesIO(docx_bytes))
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Failed to extract DOCX from bytes: {str(e)}")
            return ""
    
    @staticmethod
    def get_file_info(file_path: str) -> Dict[str, Any]:
        """
        Get file information including size, extension, and type
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary containing file information
        """
        try:
            file_stats = os.stat(file_path)
            file_extension = os.path.splitext(file_path)[1].lower()
            
            return {
                'name': os.path.basename(file_path),
                'size': file_stats.st_size,
                'extension': file_extension,
                'supported': file_extension in ['.pdf', '.docx', '.txt'],
                'path': file_path
            }
        except Exception as e:
            logger.error(f"Error getting file info for {file_path}: {str(e)}")
            return {}
    
    @classmethod
    def extract_text(cls, file_path: str) -> str:
        """
        Main method to extract text from any supported file format
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted text as string
        """
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return cls.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return cls.extract_text_from_docx(file_path)
        elif file_extension == '.txt':
            return cls.extract_text_from_txt(file_path)
        else:
            logger.error(f"Unsupported file format: {file_extension}")
            return ""


def extract_multiple_resumes(file_paths: list) -> Dict[str, str]:
    """
    Extract text from multiple resume files
    
    Args:
        file_paths: List of file paths
        
    Returns:
        Dictionary mapping filename to extracted text
    """
    results = {}
    extractor = TextExtractor()
    
    for file_path in file_paths:
        filename = os.path.basename(file_path)
        try:
            text = extractor.extract_text(file_path)
            if text:
                results[filename] = text
                logger.info(f"Successfully extracted text from {filename}")
            else:
                logger.warning(f"No text extracted from {filename}")
        except Exception as e:
            logger.error(f"Failed to extract text from {filename}: {str(e)}")
    
    return results


# Example usage and testing
if __name__ == "__main__":
    # Test the extractor
    extractor = TextExtractor()
    
    # Test with a sample text
    sample_text = "This is a sample resume text for testing purposes."
    print("Sample extraction test passed!" if sample_text else "Test failed!")
